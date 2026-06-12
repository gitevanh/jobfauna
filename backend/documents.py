"""
documents.py
------------
Turn plain text (a cover letter or a tailored CV) into a downloadable Word
(.docx) or PDF file. Stateless: it just formats whatever text it's given, so
downloads always reflect what's currently on screen (including manual edits).

Paragraphs are separated by blank lines; single newlines become line breaks
within a paragraph.
"""

import io


def _blocks(body):
    body = (body or "").replace("\r\n", "\n").replace("\r", "\n")
    out = []
    for block in body.split("\n\n"):
        block = block.strip("\n")
        if block.strip():
            out.append(block)
    return out


def build_docx(body, title=None):
    from docx import Document
    from docx.shared import Pt, Inches

    doc = Document()
    for section in doc.sections:
        section.left_margin = section.right_margin = Inches(1)
        section.top_margin = section.bottom_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    if title:
        doc.add_heading(title, level=1)

    for block in _blocks(body):
        p = doc.add_paragraph()
        for i, line in enumerate(block.split("\n")):
            if i:
                p.add_run().add_break()
            p.add_run(line)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_pdf(body, title=None):
    import html as _html
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.units import inch
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=LETTER,
        leftMargin=inch, rightMargin=inch, topMargin=inch, bottomMargin=inch,
    )
    styles = getSampleStyleSheet()
    body_style = ParagraphStyle("Body", parent=styles["Normal"],
                                fontName="Helvetica", fontSize=11, leading=15)
    title_style = ParagraphStyle("DocTitle", parent=styles["Heading1"],
                                 fontName="Helvetica-Bold", fontSize=16, spaceAfter=12)

    flow = []
    if title:
        flow.append(Paragraph(_html.escape(title), title_style))
    for block in _blocks(body):
        safe = _html.escape(block).replace("\n", "<br/>")
        flow.append(Paragraph(safe, body_style))
        flow.append(Spacer(1, 8))

    doc.build(flow)
    return buf.getvalue()


def extract_text(filename, data):
    """
    Pull plain text out of an uploaded CV. Supports PDF, DOCX, and TXT/MD.
    `data` is the raw file bytes. Raises ValueError for unsupported types.
    """
    name = (filename or "").lower()

    if name.endswith(".pdf"):
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(data))
        pages = [(page.extract_text() or "") for page in reader.pages]
        return "\n\n".join(pages).strip()

    if name.endswith(".docx"):
        from docx import Document
        doc = Document(io.BytesIO(data))
        return "\n".join(p.text for p in doc.paragraphs).strip()

    if name.endswith((".txt", ".md", ".text")):
        return data.decode("utf-8", errors="replace").strip()

    raise ValueError("Unsupported file type. Upload a PDF, DOCX, or TXT file.")

